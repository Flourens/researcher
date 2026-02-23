#!/usr/bin/env python3
"""
Fill official FFplus Part A (xlsx) and Part B (docx) templates
with data from scientific-proposal.json and organization-info.json.

Usage:
    python3 fill-official-templates.py

Output:
    results/CosTERRA-Part-A.xlsx
    results/CosTERRA-Part-B.docx
"""

import json
import os
import copy
from pathlib import Path

try:
    import openpyxl
except ImportError:
    print("Installing openpyxl...")
    os.system("pip3 install openpyxl")
    import openpyxl

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    os.system("pip3 install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
RESULTS = ROOT / "results"
TEMPLATES = RESULTS / "templates"

# ─── Load Data ───
def load_json(name: str) -> dict:
    with open(RESULTS / name, "r") as f:
        return json.load(f)

proposal = load_json("scientific-proposal.json")
org_info = load_json("organization-info.json")

# ─── Consortium Data ───
# Main participant: ALD Engineering & Construction LLC
# Supporting participant: Zaporizhzhia National University

CONSORTIUM = {
    "title": "CosTERRA",
    "full_title": "Cognitive System for Technical Engineering, Research & Architectural Documentation",
    "coordinator_name": "Evgeniy Areshkin",
    "coordinator_email": "evgeniy@ald.ua",
    "keywords": "generative AI, construction cost estimation, domain-specific LLM, EuroHPC, fine-tuning",
    "software": "PyTorch, DeepSpeed, HuggingFace Transformers, Mistral-7B/Llama-3-8B",
    "data_description": "Proprietary structured procurement database (6 years, ~15 GB, 50+ tables) from ALD's ERP system",
    "hpc_resources": "EuroHPC AI Factories pre-exascale GPU clusters (8,000 A100/H100 GPU-hours requested)",
    "num_participants": 2,
    "abstract": proposal.get("summary", proposal.get("abstract", ""))[:500],  # max ~100 words
    "participants": [
        {
            "number": 1,
            "type": "MP",
            "legal_name": "ALD Engineering & Construction LLC",
            "short_name": "ALD",
            "address": "Zaporizhzhia, Ukraine",
            "country": "UA",
            "contact_name": "Evgeniy Areshkin",
            "contact_email": "evgeniy@ald.ua",
            "telephone": "+380 XX XXX XXXX",
            "signatory": "Evgeniy Areshkin",
            "signatory_email": "evgeniy@ald.ua",
            "pic": "[TO BE OBTAINED]",
            "org_type": "SME",
            "first_eu_project": "Yes",
            "role": None,  # Not applicable for MP
        },
        {
            "number": 2,
            "type": "SP",
            "legal_name": "Zaporizhzhia National University",
            "short_name": "ZNU",
            "address": "66 Zhukovskogo St, Zaporizhzhia, 69600, Ukraine",
            "country": "UA",
            "contact_name": "Dr. Halyna Shylo",
            "contact_email": "shylo@znu.edu.ua",
            "telephone": "+380 XX XXX XXXX",
            "signatory": "Prof. [Rector Name]",
            "signatory_email": "rector@znu.edu.ua",
            "pic": "[TO BE OBTAINED]",
            "org_type": "AC",
            "first_eu_project": "No",
            "role": "HPC expert",
        },
    ],
    "budget": {
        1: {"short_name": "ALD", "rate": 1.0, "effort_pm": 18, "personnel": 120000, "other_direct": 65000, "requested": 185000},
        2: {"short_name": "ZNU", "rate": 1.0, "effort_pm": 12, "personnel": 70000, "other_direct": 30000, "requested": 100000},
    },
}


# ═══════════════════════════════════════════════════════════════
# PART A: Fill xlsx template
# ═══════════════════════════════════════════════════════════════
def fill_part_a():
    src = TEMPLATES / "Part-A-Template.xlsx"
    dst = RESULTS / "CosTERRA-Part-A.xlsx"

    wb = openpyxl.load_workbook(str(src))

    # --- Sheet: General Info ---
    ws = wb["General Info "]
    ws["B2"] = CONSORTIUM["full_title"]
    ws["B3"] = CONSORTIUM["title"]
    ws["B4"] = CONSORTIUM["coordinator_name"]
    ws["B5"] = CONSORTIUM["coordinator_email"]
    ws["B6"] = CONSORTIUM["keywords"]
    ws["B7"] = CONSORTIUM["software"]
    ws["B8"] = CONSORTIUM["data_description"]
    ws["B9"] = CONSORTIUM["hpc_resources"]
    ws["B10"] = CONSORTIUM["num_participants"]
    ws["B11"] = CONSORTIUM["abstract"]
    # Extension of Call-1 Type-2?
    ws["B12"] = "No"
    # Discussed with NCC?
    ws["B13"] = "No"
    # Discussed with AI Factory?
    ws["B14"] = "No"
    # Row 15 is a merged section header ("Information about previous similar projects") — skip it
    # Was this (or similar) proposal already submitted in FFplus Call1?
    ws["B16"] = "No"
    # Were similar or related activities funded through past Fortissimo/EC/national projects?
    ws["B17"] = "No"
    # Willing to respond to FFplus surveys?
    ws["B20"] = "Yes"
    # How did the Consortium participants become aware of the FFplus Open Call?
    ws["B23"] = "Yes"  # Electronic Media
    ws["B26"] = "Yes"  # Personal Contacts

    # --- Sheet: Participant 1 ---
    p1 = CONSORTIUM["participants"][0]
    ws1 = wb["Participant 1"]
    ws1["B4"] = p1["legal_name"]
    ws1["B5"] = p1["short_name"]
    ws1["B6"] = p1["address"]
    ws1["B7"] = p1["country"]
    ws1["B8"] = p1["contact_name"]
    ws1["B9"] = p1["contact_email"]
    ws1["B10"] = p1["telephone"]
    ws1["B11"] = p1["signatory"]
    ws1["B12"] = p1["signatory_email"]
    ws1["B13"] = p1["pic"]
    ws1["B14"] = p1["org_type"]
    # Row 15 is a note about org type definitions — skip it
    ws1["B16"] = p1["first_eu_project"]

    # --- Sheet: Participant 2 ---
    p2 = CONSORTIUM["participants"][1]
    ws2 = wb["Participant 2"]
    ws2["B4"] = p2["legal_name"]
    ws2["B5"] = p2["short_name"]
    ws2["B6"] = p2["address"]
    ws2["B7"] = p2["country"]
    ws2["B8"] = p2["contact_name"]
    ws2["B9"] = p2["contact_email"]
    ws2["B10"] = p2["telephone"]
    ws2["B11"] = p2["signatory"]
    ws2["B12"] = p2["signatory_email"]
    ws2["B13"] = p2["pic"]
    ws2["B14"] = p2["org_type"]
    # Row 15 is a note about org type definitions — skip it
    ws2["B16"] = p2["role"]
    ws2["B17"] = p2["first_eu_project"]

    # --- Sheet: Requested Funding ---
    wsf = wb["Requested Funding"]
    for pnum in [1, 2]:
        b = CONSORTIUM["budget"][pnum]
        row = 3 + pnum  # rows 4, 5
        wsf.cell(row=row, column=2, value=b["short_name"])
        wsf.cell(row=row, column=3, value=b["rate"])
        wsf.cell(row=row, column=4, value=b["effort_pm"])
        wsf.cell(row=row, column=5, value=b["personnel"])
        wsf.cell(row=row, column=6, value=b["other_direct"])
        # Column 7 (Total) is a formula, don't overwrite
        wsf.cell(row=row, column=8, value=b["requested"])

    wb.save(str(dst))
    print(f"Part A saved: {dst}")
    print(f"  Size: {dst.stat().st_size / 1024:.1f} KB")


# ═══════════════════════════════════════════════════════════════
# PART B: Fill docx tables (Work Plan + Participants/Effort)
# ═══════════════════════════════════════════════════════════════
def fill_part_b_tables(doc):
    """Fill Work Plan Table (Table 0) and Participants/Effort Table (Table 1)."""
    from copy import deepcopy

    t0 = doc.tables[0]  # Work Plan Table: 10 rows x 6 cols
    t1 = doc.tables[1]  # Participants/Effort: 3 rows x 7 cols

    def set_cell_text(table, row, col, text, font_size=10, bold=False):
        """Clear a cell and write new text."""
        cell = table.cell(row, col)
        # Clear all paragraphs
        for p in cell.paragraphs:
            p.clear()
        # Write into the first paragraph
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.bold = bold

    # ─── Table 0: Work Plan ───

    # Row 0: Innovation Study Title (merged across 6 cols)
    set_cell_text(t0, 0, 0,
        "CosTERRA \u2014 Cognitive System for Technical Engineering, "
        "Research & Architectural Documentation",
        font_size=11, bold=True)

    # Row 1: Participant short names
    set_cell_text(t0, 1, 1, "ALD")
    set_cell_text(t0, 1, 2, "ZNU")

    # Row 2: Roles
    set_cell_text(t0, 2, 1, "SME (Main Participant)")
    set_cell_text(t0, 2, 2, "HPC Expert (Supporting)")

    # Row 3: Description (merged across 6 cols)
    set_cell_text(t0, 3, 0,
        "CosTERRA develops a domain-specific generative AI model (7\u201313B parameters) "
        "for automated generation of construction cost estimation documents compliant "
        "with Ukrainian building standards (DBN, DSTU). The model is fine-tuned on "
        "ALD\u2019s proprietary dataset of \u226550,000 instruction\u2013output pairs using "
        "EuroHPC pre-exascale GPU infrastructure, with a custom structured-data "
        "encoder for real-time ERP database integration.",
        font_size=10)

    # Tasks data
    tasks = [
        {
            "name": "Data Engineering",
            "duration": "M1\u2013M3",
            "participants": "ALD: 6 PM, ZNU: 1 PM",
            "deliverable": "D1.1: Validated training corpus (M3)",
            "description": (
                "Extract structured data from ALD ERP (PostgreSQL); transform into "
                "instruction\u2013output pairs; augment with standards references; validate "
                "quality (automated + expert review); establish data versioning with DVC."
            ),
            "compute": "ALD development servers; 100 GPU-hours for data quality profiling",
        },
        {
            "name": "Architecture Design & Baselines",
            "duration": "M2\u2013M4",
            "participants": "ZNU: 3 PM, ALD: 2 PM",
            "deliverable": "D2.1: Model architecture specification (M3); D2.2: Baseline evaluation report (M4)",
            "description": (
                "Design structured-data encoder; implement integration with transformer "
                "base model; benchmark 4 baseline systems on held-out test set; select "
                "optimal base model; implement distributed training pipeline (FSDP)."
            ),
            "compute": "500 GPU-node-hours for baseline experiments",
        },
        {
            "name": "Model Training & Optimisation",
            "duration": "M4\u2013M8",
            "participants": "ZNU: 4 PM, ALD: 3 PM",
            "deliverable": "D3.1: Trained CosTERRA model (M8)",
            "description": (
                "Full-parameter fine-tuning on EuroHPC; 3+ training runs with "
                "hyperparameter search (learning rate, batch size, sequence length); "
                "mixed-precision training; checkpoint management; convergence monitoring "
                "via MLflow."
            ),
            "compute": "6,400 GPU-node-hours (main training)",
        },
        {
            "name": "Evaluation & Benchmarking",
            "duration": "M7\u2013M9",
            "participants": "ALD: 4 PM, ZNU: 2 PM",
            "deliverable": "D4.1: Month 7 pre-final report (M7); D4.2: Final evaluation report (M9)",
            "description": (
                "Evaluate against all metrics (compliance, numerical consistency, "
                "ROUGE-L, BERTScore, expert acceptance); ablation studies; comparison "
                "against baselines; error analysis and failure mode documentation."
            ),
            "compute": "800 GPU-node-hours for evaluation",
        },
        {
            "name": "Integration Design & Reporting",
            "duration": "M8\u2013M10",
            "participants": "ALD: 5 PM, ZNU: 1 PM",
            "deliverable": "D5.1: Integration architecture (M9); D5.2: Final report (M10); D5.3: Success story (M10)",
            "description": (
                "Design API interface for ERP integration (architecture only, no "
                "deployment); document model card; prepare final report with business "
                "impact analysis; contribute to FFplus success story."
            ),
            "compute": "200 GPU-node-hours for demonstration",
        },
    ]

    def format_task(idx, task):
        return (
            f"Task {idx}: {task['name']}\n"
            f"Duration: {task['duration']}\n"
            f"Participants: {task['participants']}\n"
            f"Deliverable: {task['deliverable']}\n\n"
            f"Technical description:\n{task['description']}\n\n"
            f"Computational Resources: {task['compute']}"
        )

    # Fill existing task rows 4, 5, 6 with Tasks 1-3
    for i in range(3):
        set_cell_text(t0, 4 + i, 0, format_task(i + 1, tasks[i]), font_size=10)

    # Insert 2 new rows after row 6 for Tasks 4-5
    # Copy the XML structure of row 6 (a merged task row) and insert after it
    row6_tr = t0.rows[6]._tr
    for i in range(2):
        new_tr = deepcopy(row6_tr)
        # Clear text in the new row
        for tc in new_tr.iterchildren():
            for p in tc.iterchildren():
                if p.tag.endswith('}p'):
                    for r in list(p):
                        if r.tag.endswith('}r'):
                            p.remove(r)
        row6_tr.addnext(new_tr)
        row6_tr = new_tr  # next insertion goes after this one

    # After insertion: rows 7-8 are new task rows, old rows 7-9 shifted to 9-11
    # Fill tasks 4-5 in the new rows (now rows 7, 8)
    set_cell_text(t0, 7, 0, format_task(4, tasks[3]), font_size=10)
    set_cell_text(t0, 8, 0, format_task(5, tasks[4]), font_size=10)

    # Row 10 (was 8): Computational Resources (merged)
    set_cell_text(t0, 10, 0,
        "Total required GPU-node-hours: 8,000 on A100/H100-class GPU nodes "
        "(8 GPUs per node, equivalent to 64,000 GPU-hours).\n\n"
        "EuroHPC access plan: The consortium will apply for EuroHPC AI Factories "
        "access. Target systems: LUMI (CSC, Finland), MareNostrum 5 (BSC, Spain), "
        "or Leonardo (CINECA, Italy). Application will be submitted with HPC "
        "National Competence Centre assistance.\n\n"
        "Software: PyTorch \u22652.2, HuggingFace Transformers, DeepSpeed/FSDP, MLflow, "
        "CUDA \u226512.0. Storage: ~10 TB for training data, model checkpoints, "
        "and experiment logs.",
        font_size=10)

    # Row 11 (was 9): Impact and Outputs (merged)
    set_cell_text(t0, 11, 0,
        "Outputs:\n"
        "\u2022 Validated training corpus of \u226550,000 instruction\u2013output pairs "
        "for construction cost estimation\n"
        "\u2022 CosTERRA generative model: fine-tuned 7\u201313B parameter transformer "
        "with structured-data encoder\n"
        "\u2022 Comprehensive evaluation report with metric scores against 4 baselines\n"
        "\u2022 Model card and integration architecture specification\n"
        "\u2022 Open-source evaluation benchmark for construction document generation\n"
        "\u2022 Data management plan and trustworthy AI assessment\n\n"
        "Impact:\n"
        "\u2022 Validated technical feasibility for AI-powered cost estimation "
        "(de-risks \u20ac1.5M+ annual savings)\n"
        "\u2022 Production-ready model architecture for integration into ALD\u2019s SaaS platform\n"
        "\u2022 Competitive advantage in Ukrainian construction technology market\n"
        "\u2022 Foundation for regional market expansion (Kazakhstan, Georgia, Moldova)\n"
        "\u2022 Demonstrated EuroHPC utilisation pathway for SME generative AI development",
        font_size=10)

    print("  Filled Work Plan Table (5 tasks + resources + impact)")

    # ─── Table 1: Participants and Effort ───
    set_cell_text(t1, 1, 1, "ALD")
    set_cell_text(t1, 1, 2, "ZNU")
    set_cell_text(t1, 2, 1, "20")
    set_cell_text(t1, 2, 2, "11")
    set_cell_text(t1, 2, 6, "31")

    print("  Filled Participants and Effort Table (ALD: 20 PM, ZNU: 11 PM, Total: 31 PM)")


# ═══════════════════════════════════════════════════════════════
# PART B: Fill docx template
# ═══════════════════════════════════════════════════════════════
def fill_part_b():
    src = TEMPLATES / "Part-B-Template.docx"
    dst = RESULTS / "CosTERRA-Part-B.docx"

    doc = Document(str(src))

    # Helper: Find and replace text in paragraphs
    def set_paragraph_text(paragraph, new_text: str, font_name="Arial", font_size=11, bold=False):
        """Clear paragraph and set new text while preserving paragraph style."""
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            run = paragraph.runs[0]
        else:
            run = paragraph.add_run()
        run.text = new_text
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold

    # Helper: Add formatted content from markdown-like text
    def add_section_content(doc, text: str, insert_before_idx: int = None):
        """Parse markdown-like text and add as formatted paragraphs."""
        lines = text.split("\n")
        new_paragraphs = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Skip top-level section headers (## )
            if stripped.startswith("## "):
                continue

            # Subsection headers (### )
            if stripped.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[4:])
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True
                p.space_after = Pt(4)
                p.space_before = Pt(8)
                new_paragraphs.append(p)
                continue

            # Table lines — add as-is for now (tables need special handling in docx)
            if stripped.startswith("|") and stripped.endswith("|"):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.name = "Arial"
                run.font.size = Pt(9)
                new_paragraphs.append(p)
                continue

            # Bullet points
            if stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Paragraph")
                text_content = stripped[2:]
                # Handle bold markers
                parts = text_content.split("**")
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold part
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                new_paragraphs.append(p)
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            # Handle bold markers
            parts = stripped.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold part
                    run = p.add_run(part)
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = "Arial"
                run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            new_paragraphs.append(p)

        return new_paragraphs

    # ─── Strategy: Replace template content ───
    # The docx template has placeholder/instruction text that we need to replace.
    # We'll find section headers and replace the content between them.

    # First, fill the cover page fields
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "Innovation Study Title":
            set_paragraph_text(para, f"CosTERRA — {CONSORTIUM['full_title']}", bold=True, font_size=14)
        elif text.startswith("Name of the coordinating person:"):
            set_paragraph_text(para, f"Name of the coordinating person: {CONSORTIUM['coordinator_name']}, ALD Engineering & Construction LLC")
        elif text.startswith("E-mail:"):
            set_paragraph_text(para, f"E-mail: {CONSORTIUM['coordinator_email']}")

    # Now find section boundaries and replace content
    # Sections to find:
    sections_map = {
        "Summary": proposal.get("summary", proposal.get("abstract", "")),
        "Industrial relevance, potential impact and exploitation plans": proposal.get("industrialRelevance", ""),
        "Description of the work plan, technological/algorithmic approach and software development strategy": proposal.get("workPlan", ""),
        "Quality of the consortium as a whole and of the individual proposers": proposal.get("consortiumQuality", ""),
        "Justification of costs and resources": proposal.get("costJustification", ""),
    }

    # Find paragraph indices for each section header
    section_indices = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for section_title in sections_map:
            if text.startswith(section_title):
                section_indices[section_title] = i
                break

    print(f"  Found {len(section_indices)} section headers in template")

    # For each section, remove italic instruction text and add our content
    # We'll work backwards to preserve indices
    section_titles_ordered = list(sections_map.keys())

    for section_title in section_titles_ordered:
        if section_title not in section_indices:
            print(f"  WARNING: Section '{section_title}' not found in template!")
            continue

        idx = section_indices[section_title]
        content = sections_map[section_title]
        if not content:
            print(f"  WARNING: No content for section '{section_title}'")
            continue

        # Find the next section header index (or end of document)
        next_idx = len(doc.paragraphs)
        for other_title, other_idx in section_indices.items():
            if other_idx > idx and other_idx < next_idx:
                next_idx = other_idx

        # Remove all paragraphs between this header and the next (italic instructions)
        # We can't easily delete paragraphs from python-docx, so we clear their text
        # and add our content after the header
        for j in range(idx + 1, next_idx):
            para = doc.paragraphs[j]
            # Clear the instruction text
            for run in para.runs:
                run.text = ""

        # Now add our content right after the section header
        # Since we can't insert paragraphs at specific positions easily,
        # we'll rewrite the first cleared paragraph with our content
        content_lines = content.split("\n")
        para_idx = idx + 1
        for line in content_lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                continue

            if para_idx < next_idx:
                # Reuse existing (cleared) paragraph
                para = doc.paragraphs[para_idx]
                para.clear()

                if stripped.startswith("### "):
                    run = para.add_run(stripped[4:])
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    run.bold = True
                    para.paragraph_format.space_before = Pt(8)
                    para.paragraph_format.space_after = Pt(4)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    text_content = stripped[2:]
                    parts = text_content.split("**")
                    for k, part in enumerate(parts):
                        run = para.add_run(part)
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        if k % 2 == 1:
                            run.bold = True
                elif stripped.startswith("|") and stripped.endswith("|"):
                    run = para.add_run(stripped)
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                else:
                    parts = stripped.split("**")
                    for k, part in enumerate(parts):
                        run = para.add_run(part)
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        if k % 2 == 1:
                            run.bold = True
                    para.paragraph_format.space_after = Pt(3)

                para_idx += 1
            else:
                # We've run out of placeholder paragraphs, need to insert new ones
                # python-docx doesn't support easy insertion, so we append at the end
                # and note this limitation
                break

    # Add references at the end
    if proposal.get("bibliography"):
        doc.add_paragraph()  # spacing
        p = doc.add_paragraph()
        run = p.add_run("References")
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.bold = True

        for ref in proposal["bibliography"]:
            p = doc.add_paragraph()
            run = p.add_run(ref)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)

    # ─── Fill Tables ───
    fill_part_b_tables(doc)

    doc.save(str(dst))
    print(f"Part B saved: {dst}")
    print(f"  Size: {dst.stat().st_size / 1024:.1f} KB")
    print(f"  NOTE: Review in Word — check formatting, embedded Excel cost table, and page count")
    print(f"  NOTE: Double-click the embedded cost table in Word to fill budget data")


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 60)
    print("  FFplus Template Filler — CosTERRA")
    print("=" * 60)
    print()

    # Check inputs exist
    for f in ["scientific-proposal.json", "organization-info.json"]:
        if not (RESULTS / f).exists():
            print(f"ERROR: {f} not found in results/")
            exit(1)

    # Check templates exist
    for f in ["Part-A-Template.xlsx", "Part-B-Template.docx"]:
        if not (TEMPLATES / f).exists():
            print(f"ERROR: Template {f} not found in results/templates/")
            exit(1)

    schema = "FFplus Part B (5 sections)" if "summary" in proposal else "Legacy (7 sections)"
    print(f"Proposal schema: {schema}")
    print()

    print("--- Filling Part A (xlsx) ---")
    fill_part_a()
    print()

    print("--- Filling Part B (docx) ---")
    fill_part_b()
    print()

    print("=" * 60)
    print("  Done! Next steps:")
    print("  1. Open CosTERRA-Part-A.xlsx in Excel — add PIC codes, verify data")
    print("  2. Open CosTERRA-Part-B.docx in Word — check formatting, fill embedded Excel")
    print("  3. Export Part B as PDF (max 5 MB, max 13 pages excl. cover)")
    print("  4. Submit both files via the FFplus submission portal")
    print("=" * 60)
