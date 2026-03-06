#!/usr/bin/env python3
"""
Fill the Nansen EDU 2025 Project Description DOCX template
for all Smart Grid variants (EN + UA where available).

Converts markdown formatting from AI output into proper DOCX formatting:
- **bold** → bold runs
- *italic* → italic runs
- ### / ## headers → bold paragraphs
- --- → removed
- - bullet → indented paragraph with dash
- Proper paragraph spacing and indentation
"""

import json
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, 'templates', 'project-description.docx')
VARIANTS_DIR = os.path.join(SCRIPT_DIR, 'smartgrid-variants')

VARIANT_META = {
    'variant-a': {
        'title': "AI-Powered Smart Grid Education for Ukraine's Energy Resilience",
        'acronym': 'AI-SmartGrid',
    },
    'variant-b': {
        'title': "Smart and Resilient Energy Systems for Ukraine's Post-War Reconstruction",
        'acronym': 'SmartEnergy-UA',
    },
    'variant-c': {
        'title': "Digital Twins and AI for Flexible Energy Systems Education",
        'acronym': 'DigiTwin-Energy',
    },
    'variant-d': {
        'title': "Smart Grid and Green Energy Transition Education for Ukraine",
        'acronym': 'SmartGreen-UA',
    },
}

FONT_NAME = 'Arial'
FONT_SIZE = Pt(11)
FONT_SIZE_SMALL = Pt(10)
FONT_COLOR = RGBColor(0, 0, 0)


def add_run(paragraph, text, bold=False, italic=False, font_size=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size or FONT_SIZE
    run.font.color.rgb = FONT_COLOR
    run.bold = bold
    run.italic = italic
    return run


def parse_inline_formatting(paragraph, text, font_size=None):
    """Parse **bold** and *italic* markers and add as formatted runs."""
    # Pattern matches: **bold text** or *italic text*
    # Process bold first, then italic
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(paragraph, part[2:-2], bold=True, font_size=font_size)
        elif part.startswith('*') and part.endswith('*'):
            add_run(paragraph, part[1:-1], italic=True, font_size=font_size)
        else:
            add_run(paragraph, part, font_size=font_size)


def classify_line(line):
    """Classify a line by its markdown type."""
    stripped = line.strip()
    if not stripped:
        return 'empty', ''
    if stripped == '---':
        return 'separator', ''
    if stripped.startswith('#### '):
        return 'heading4', stripped[5:]
    if stripped.startswith('### '):
        return 'heading3', stripped[4:]
    if stripped.startswith('## '):
        return 'heading2', stripped[3:]
    if stripped.startswith('# '):
        return 'heading1', stripped[2:]
    if re.match(r'^[-•]\s', stripped):
        return 'bullet', stripped[2:]
    if re.match(r'^\d+[\.\)]\s', stripped):
        # Numbered list: "1. text" or "1) text"
        m = re.match(r'^(\d+[\.\)])\s(.+)', stripped)
        return 'numbered', stripped
    if stripped.startswith('|') and '|' in stripped[1:]:
        return 'table_row', stripped
    return 'text', stripped


def add_formatted_paragraph(doc, line_type, text, after_element):
    """Create a properly formatted paragraph and insert after the given element."""
    new_para = doc.add_paragraph()
    new_para.style = doc.styles['Normal']

    pf = new_para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

    if line_type in ('heading1', 'heading2', 'heading3', 'heading4'):
        # Bold subheading, slightly larger spacing
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        # Remove section number duplicates like "1.1 Background..."
        # since section headers are already in template
        clean = re.sub(r'^[0-9]+\.[0-9]+\s*', '', text).strip()
        if clean:
            text = clean
        add_run(new_para, text, bold=True)

    elif line_type == 'bullet':
        # Indent bullet points
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.35)
        new_para.add_run('– ')  # Use dash instead of markdown bullet
        parse_inline_formatting(new_para, text)

    elif line_type == 'numbered':
        pf.left_indent = Cm(0.5)
        parse_inline_formatting(new_para, text)

    elif line_type == 'table_row':
        # Convert markdown table to tab-separated text
        cells = [c.strip() for c in text.split('|') if c.strip()]
        if all(re.match(r'^[-:]+$', c) for c in cells):
            # This is a separator row (|---|---|), skip it
            new_para._element.getparent().remove(new_para._element)
            return after_element
        tab_text = '\t'.join(cells)
        add_run(new_para, tab_text)

    else:
        # Regular paragraph
        pf.space_after = Pt(6)
        parse_inline_formatting(new_para, text)

    after_element.addnext(new_para._element)
    return new_para._element


def process_section_content(doc, content, after_element):
    """
    Process section content: split into lines, classify each,
    and create properly formatted DOCX paragraphs.
    """
    lines = content.split('\n')
    current_element = after_element

    # Group consecutive lines into logical blocks
    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, text = classify_line(line)

        if line_type == 'empty':
            i += 1
            continue

        if line_type == 'separator':
            # Skip markdown horizontal rules
            i += 1
            continue

        # For regular text, merge consecutive non-empty lines that aren't
        # special types (to form proper paragraphs)
        if line_type == 'text':
            # Check if next lines continue this paragraph (no blank line between)
            merged = text
            while i + 1 < len(lines):
                next_type, next_text = classify_line(lines[i + 1])
                if next_type == 'text':
                    merged += ' ' + next_text
                    i += 1
                else:
                    break
            current_element = add_formatted_paragraph(
                doc, 'text', merged, current_element
            )
        else:
            current_element = add_formatted_paragraph(
                doc, line_type, text, current_element
            )

        i += 1

    return current_element


def fill_template(proposal, variant_id, lang, output_path):
    """Fill the DOCX template with proposal content."""
    doc = Document(TEMPLATE_PATH)
    meta = VARIANT_META[variant_id]

    # --- Fill Cover Table ---
    cover_table = doc.tables[0]
    cover_table.rows[1].cells[1].text = f"{meta['title']} ({meta['acronym']})"
    cover_table.rows[2].cells[1].text = "NTNU — Norwegian University of Science and Technology"

    # --- Delete guideline tables (Tables 1–14) ---
    for idx in reversed(range(1, min(15, len(doc.tables)))):
        table = doc.tables[idx]
        table._element.getparent().remove(table._element)

    # --- Map sections to "► Delete..." markers ---
    section_map = {
        '1.1': proposal['section1_1_background'],
        '1.2': proposal['section1_2_alignment'],
        '2.1': proposal['section2_1_activities'],
        '2.2': proposal['section2_2_risks'],
        '3.1': proposal['section3_1_personnel'],
        '3.2': proposal['section3_2_collaboration'],
        '4.1': proposal['section4_1_impact'],
        '4.2': proposal['section4_2_sustainability'],
    }

    current_section = None
    for para in doc.paragraphs:
        text = para.text.strip()

        if text.startswith('1.1'):
            current_section = '1.1'
        elif text.startswith('1.2'):
            current_section = '1.2'
        elif text.startswith('2.1'):
            current_section = '2.1'
        elif text.startswith('2.2'):
            current_section = '2.2'
        elif text.startswith('3.1'):
            current_section = '3.1'
        elif text.startswith('3.2'):
            current_section = '3.2'
        elif text.startswith('4.1'):
            current_section = '4.1'
        elif text.startswith('4.2'):
            current_section = '4.2'
        elif '► Delete' in text and current_section in section_map:
            content = section_map[current_section]

            # Clear the marker paragraph — use it as anchor
            para.text = ''

            # Process the section content with proper formatting
            process_section_content(doc, content, para._element)

            current_section = None

    # --- Fill References ---
    references = proposal.get('references', [])
    ref_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in ('[Reference]', '[...]'):
            ref_indices.append(i)

    if ref_indices and references:
        first_ref_para = doc.paragraphs[ref_indices[0]]
        first_ref_para.text = ''
        # Clean markdown from references too
        clean_ref = re.sub(r'\*\*([^*]+)\*\*', r'\1', references[0])
        clean_ref = re.sub(r'\*([^*]+)\*', r'\1', clean_ref)
        add_run(first_ref_para, clean_ref, font_size=FONT_SIZE_SMALL)

        current_element = first_ref_para._element
        for ref in references[1:]:
            new_para = doc.add_paragraph()
            new_para.style = doc.styles['Normal']
            clean_ref = re.sub(r'\*\*([^*]+)\*\*', r'\1', ref)
            clean_ref = re.sub(r'\*([^*]+)\*', r'\1', clean_ref)
            add_run(new_para, clean_ref, font_size=FONT_SIZE_SMALL)
            current_element.addnext(new_para._element)
            current_element = new_para._element

        # Remove old placeholders
        for idx in ref_indices[1:]:
            para = doc.paragraphs[idx]
            if para.text.strip() in ('[Reference]', '[...]'):
                para._element.getparent().remove(para._element)

    # Clean remaining placeholders
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t in ('[Reference]', '[...]'):
            para._element.getparent().remove(para._element)
        elif 'List of references referred to in the project description' in t:
            para.text = ''

    # Remove empty "► Delete" anchor paragraphs
    for para in list(doc.paragraphs):
        if para.text.strip() == '' and para.style.name == 'No Spacing':
            # Only remove if it's one of the old markers (now empty)
            pass  # Keep empty lines for spacing

    doc.save(output_path)
    return output_path


def main():
    print('=' * 60)
    print('FILLING DOCX TEMPLATES — PROPER FORMATTING')
    print('=' * 60)

    results = []

    for variant_id, meta in VARIANT_META.items():
        variant_dir = os.path.join(VARIANTS_DIR, variant_id)
        if not os.path.exists(variant_dir):
            print(f'\n  ✗ {meta["acronym"]}: directory not found')
            continue

        print(f'\n--- {meta["acronym"]} ---')

        # English
        en_json = os.path.join(variant_dir, 'project-description-en.json')
        if os.path.exists(en_json):
            proposal = json.load(open(en_json, 'r', encoding='utf-8'))
            out = os.path.join(variant_dir, f'{meta["acronym"]}-Project-Description-EN.docx')
            fill_template(proposal, variant_id, 'en', out)
            words = sum(len(proposal[k].split()) for k in proposal if k != 'references' and isinstance(proposal[k], str))
            print(f'  ✓ EN: {os.path.basename(out)} (~{words} words)')
            results.append((meta['acronym'], 'EN', words, out))

        # Ukrainian
        uk_json = os.path.join(variant_dir, 'project-description-uk.json')
        if os.path.exists(uk_json):
            proposal = json.load(open(uk_json, 'r', encoding='utf-8'))
            out = os.path.join(variant_dir, f'{meta["acronym"]}-Project-Description-UA.docx')
            fill_template(proposal, variant_id, 'uk', out)
            words = sum(len(proposal[k].split()) for k in proposal if k != 'references' and isinstance(proposal[k], str))
            print(f'  ✓ UA: {os.path.basename(out)} (~{words} words)')
            results.append((meta['acronym'], 'UA', words, out))
        else:
            print(f'  — UA: not available')

    # Summary
    print(f'\n{"=" * 60}')
    print('SUMMARY')
    print(f'{"=" * 60}')
    for acronym, lang, words, fpath in results:
        print(f'  {acronym:<20} {lang}  ~{words} words  {os.path.basename(fpath)}')

    # Review scores
    print(f'\n{"=" * 60}')
    print('REVIEW SCORES')
    print(f'{"=" * 60}')
    for variant_id, meta in VARIANT_META.items():
        rp = os.path.join(VARIANTS_DIR, variant_id, 'review-results.json')
        if os.path.exists(rp):
            r = json.load(open(rp, 'r', encoding='utf-8'))
            scores = ', '.join(f'{cs["score"]}/7' for cs in r.get('criterionScores', []))
            q = 'YES' if r.get('qualifiesForFunding') else 'NO'
            print(f'  {meta["acronym"]:<20} {r["overallScore"]}/7  Qualifies: {q}  [{scores}]')

    print(f'\nDone! Files in: {VARIANTS_DIR}/')


if __name__ == '__main__':
    main()
