#!/usr/bin/env python3
"""
Fill the Nansen EDU 2025 Project Description template with AI-GreenMet content.
Project: AI-Driven Green Metallurgy Education for Ukraine's Industrial Reconstruction (AI-GreenMet)
"""

import copy
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, 'templates', 'project-description.docx')
OUTPUT_PATH = os.path.join(SCRIPT_DIR, 'AI-GreenMet-Project-Description.docx')

# ============================================================
# SECTION CONTENT
# ============================================================

SECTION_1_1 = """Ukraine's metallurgical sector, historically contributing 12–14% of GDP and over 30% of export revenue, faces an unprecedented crisis. The full-scale Russian invasion that began in February 2022 has caused estimated damages exceeding USD 10 billion to industrial infrastructure [1], with the Zaporizhzhia region — home to one of Europe's largest metallurgical clusters — suffering particularly severe destruction. The Zaporizhstal steelworks, Motor Sich engine manufacturing complex, and dozens of smaller metallurgical enterprises have been damaged or forced to reduce operations. Simultaneously, the global steel industry is undergoing a green transition driven by the European Green Deal and Carbon Border Adjustment Mechanism (CBAM), requiring decarbonisation of production processes through electrification, hydrogen-based reduction, and AI-driven process optimisation [2].

This dual challenge — war-driven destruction and the imperative of green transition — creates an urgent need for a new generation of engineers who can combine traditional metallurgical knowledge with artificial intelligence and machine learning (AI/ML) competencies. Currently, no Ukrainian higher education institution offers an integrated curriculum combining AI/ML with metallurgical engineering. This gap threatens Ukraine's ability to reconstruct its industrial base according to modern European environmental standards and to maintain competitiveness in global metal markets.

The project "AI-Driven Green Metallurgy Education for Ukraine's Industrial Reconstruction" (AI-GreenMet) directly addresses this gap by developing, piloting, and institutionalising a joint Norwegian-Ukrainian educational programme that equips students and academic staff with interdisciplinary competencies at the intersection of AI, green metallurgy, and industrial digitalisation.

Target groups and their needs:

(a) Engineering students at Ukrainian HEIs (primary target: 150+ students over 3 years). These students currently receive metallurgical education based on Soviet-era curricula with limited exposure to AI/ML methods, digital twins, or green production technologies. They need modern interdisciplinary skills to be employable in a reconstructed, EU-aligned metallurgical industry.

(b) Academic staff at Ukrainian HEIs (target: 25+ staff members). Faculty in metallurgy, materials science, and computer science departments lack training in cross-disciplinary pedagogy and access to state-of-the-art AI tools for metallurgical applications. They need professional development to deliver the new curriculum.

(c) Metallurgical industry professionals and technicians (target: 50+ through short courses and Living Lab activities). The Zaporizhzhia metallurgical cluster requires upskilled workers who can operate AI-assisted production systems, digital twins, and green metallurgy processes during reconstruction.

(d) Vocational students at ZNU's Metallurgical Professional College (target: 40+ students). These students need practical AI-enhanced skills applicable to technician-level roles in reconstructed metallurgical facilities.

Baseline data: ZNU's Engineering Institute currently offers 6 metallurgy-related degree programmes with approximately 400 enrolled students. None include dedicated AI/ML modules. ZNU's Faculty of Mathematics has an active Laboratory of Parallel and Distributed Computing with expertise in AI model development, but no joint courses with the Engineering Institute exist. NTNU's Department of Materials Science and Engineering ranks among the world's top 50 in metallurgy and materials science (QS 2024) and operates multiple SFI centres in metal production with extensive industry partnerships [3].

Planned results (SMART objectives):

1. Develop 4 new interdisciplinary course modules: (i) AI/ML Fundamentals for Metallurgical Engineers, (ii) Digital Twins for Metallurgical Processes, (iii) Green Metallurgy and Decarbonisation Technologies, (iv) Data-Driven Quality Control in Metal Production. Target: modules validated and integrated into ZNU curricula by Month 24.

2. Train 25+ ZNU academic staff through NTNU-led workshops and exchange visits, with each participant completing a minimum 40-hour professional development programme. Target: training completed by Month 30.

3. Establish a Living Laboratory for AI-Enhanced Metallurgy at ZNU, connecting the Engineering Institute's materials laboratories with the Faculty of Mathematics' computing infrastructure. Target: operational by Month 18.

4. Deliver pilot courses to 150+ students across 3 cohorts (Years 1–3), with pre/post competency assessments demonstrating measurable learning gains.

5. Produce 5+ joint scientific publications and 4 open educational resources (OER) shared via Ukrainian and Norwegian educational repositories.

6. Establish a sustainable partnership framework enabling continued cooperation through co-supervised PhDs, joint Horizon Europe proposals, and industry engagement beyond the project period.

These objectives directly support all three Nansen EDU programme goals: (1) strengthening Ukrainian HEI capacity through new curricula and staff development, (2) improving access to qualified professionals by training engineers with AI/metallurgy competencies urgently needed for reconstruction, and (3) establishing long-term Norway-Ukraine cooperation through the NTNU-ZNU institutional partnership and joint roadmap extending to 2028 and beyond [4]."""

SECTION_1_2 = """The AI-GreenMet project builds strategically on existing activities and expertise of both partner institutions, while introducing genuinely innovative elements that distinguish it from current initiatives.

Building on partners' existing activities:

ZNU has developed significant research capacity in metallurgy-related areas, including: design and fabrication of high-entropy alloys and intermetallic compounds under thermochemical pressing; engineering of electrolyser systems for magnesium extraction from fluoride-oxide melts for Al-Mg alloy production; and advancement of catalytic technologies for industrial gas purification from CO and hydrocarbons [5]. ZNU's Faculty of Mathematics operates the Laboratory of Parallel and Distributed Computing with active projects in AI for logistics and national security applications (funded at UAH 3.59 million by the Ukrainian state budget). The university also has ongoing work in smart grid technologies, digital twins for industrial systems, and renewable energy. AI-GreenMet will connect these currently separate streams — metallurgical research and AI/computing expertise — into a coherent educational programme.

NTNU brings world-class expertise through its SFI Metal Production Centre (2015–2023) and successor initiatives, which have established industry-academic partnerships with major Nordic metal producers. NTNU's Department of Materials Science and Engineering has pioneered the application of machine learning to alloy design, process optimisation, and predictive maintenance in metallurgical production. The university's experience with the Norwegian Centres of Excellence model and extensive Horizon Europe track record provide a proven framework for international research-education cooperation [3].

Knowledge of existing initiatives and added value:

Several European programmes support Ukrainian higher education, but none specifically target the AI-metallurgy intersection: (a) Erasmus+ CBHE projects in Ukraine focus on general institutional capacity building, governance reform, and quality assurance — not discipline-specific technical curricula; (b) the DAAD Ukraine Digital programme, in which ZNU participates, focuses on general digital competencies rather than AI for specific industrial sectors; (c) Horizon MSCA actions support individual researcher mobility but do not fund curriculum development or institutional capacity building at scale; (d) UNIDO and EU4Environment programmes address industrial greening but lack educational components [6]. AI-GreenMet fills a clear gap by being the first programme to systematically integrate AI/ML training into metallurgical engineering education in Ukraine.

Innovative elements:

(1) First AI-Metallurgy curriculum in Ukraine. No Ukrainian HEI currently offers integrated courses combining AI/ML with metallurgical engineering. AI-GreenMet will create and validate a modular curriculum that can be replicated by other Ukrainian technical universities.

(2) Living Laboratory approach. The project will establish a physical-digital Living Lab at ZNU that connects metallurgical laboratory equipment with AI/ML computing infrastructure. Students will work with real metallurgical datasets (alloy composition, process parameters, quality metrics) from ZNU's research projects and industry partners, using AI tools to optimise processes and predict outcomes. This experiential learning approach goes beyond traditional lecture-based education [7].

(3) Digital Twin pedagogy. Building on ZNU's existing digital twin research, the project will develop educational digital twins of metallurgical processes (e.g., electrolyser operation, alloy casting) that students can manipulate, experiment with, and optimise using AI algorithms in a safe virtual environment. This approach is novel in the context of Ukrainian metallurgical education.

(4) Blended delivery model adapted for conflict context. Recognising Zaporizhzhia's frontline status, the project has designed a hybrid delivery model from the outset — not as an emergency fallback but as a core pedagogical feature. This includes asynchronous online modules (accessible during security alerts), intensive in-person laboratory sessions (scheduled during stable periods), and remote access to NTNU's computing resources for AI model training.

(5) Industry-connected curriculum. Through ZNU's existing connections to the Zaporizhzhia metallurgical cluster (Zaporizhstal, regional SMEs) and NTNU's SFI industry partnerships, the curriculum will incorporate real industry problems, guest lectures from practitioners, and student industry projects — ensuring graduates meet actual employer needs during reconstruction [8]."""

SECTION_2_1 = """The project is organised into five Work Packages (WPs), implemented over three years (July 2026 – June 2029) aligned with the ZNU-NTNU roadmap phases: Foundation (Year 1), Pilots (Year 2), and Scaling (Year 3).

WP1: Joint Curriculum Development (Months 1–24)
Lead: NTNU | Co-lead: ZNU Faculty of Mathematics & Engineering Institute

Activities:
- M1–M6: Conduct needs assessment with Ukrainian industry stakeholders (Zaporizhstal, regional employers) and comparative analysis of European AI-metallurgy programmes. Establish curriculum working group with representatives from NTNU, ZNU Engineering Institute, ZNU Faculty of Mathematics, and ZNU Metallurgical College.
- M4–M12: Develop 4 course module syllabi with learning outcomes aligned to the European Qualifications Framework (EQF levels 6–7): (i) AI/ML Fundamentals for Metallurgical Engineers (6 ECTS), (ii) Digital Twins for Metallurgical Processes (5 ECTS), (iii) Green Metallurgy and Decarbonisation Technologies (5 ECTS), (iv) Data-Driven Quality Control in Metal Production (4 ECTS).
- M10–M18: Develop teaching materials: lecture slides, laboratory exercises, programming assignments (Python, TensorFlow/PyTorch for metallurgical data), case studies from NTNU SFI research and ZNU industry data.
- M18–M24: External review of curricula by NTNU quality assurance unit and Ukrainian Ministry of Education standards. Integration into ZNU's official degree programme catalogue.

Deliverables: D1.1 Needs assessment report (M6); D1.2 Four module syllabi and teaching plans (M12); D1.3 Complete teaching materials package (M18); D1.4 Accredited curriculum modules in ZNU system (M24).

WP2: Staff Training and Exchanges (Months 3–33)
Lead: NTNU | Co-lead: ZNU Engineering Institute

Activities:
- M3–M6: Select 25 ZNU staff (from Engineering Institute, Faculty of Mathematics, and Metallurgical College) for training programme. Selection criteria: disciplinary relevance, teaching load, English proficiency, commitment to curriculum implementation.
- M6–M12: NTNU-led online training series (8 workshops × 4 hours) covering: AI/ML for metallurgy, digital twin development, problem-based learning pedagogy, open educational resource creation.
- M9–M15: First round of staff exchanges: 8 ZNU staff visit NTNU (up to 30 days each) for intensive training in NTNU laboratories, observation of teaching methods, and joint curriculum development sessions.
- M15–M24: Second round: 8 ZNU staff exchanges to NTNU. NTNU staff (3–4 persons) visit ZNU for co-teaching and Living Lab setup support.
- M24–M33: Third round: 5 ZNU staff exchanges focused on advanced topics and co-supervision of student projects. Mentoring of junior ZNU staff by trained colleagues.

Deliverables: D2.1 Staff training programme design (M6); D2.2 Training completion certificates for 25 staff (M30); D2.3 Staff exchange reports with knowledge transfer documentation (M12, M24, M33).

WP3: Student Training Programme (Months 10–36)
Lead: ZNU Engineering Institute | Co-lead: ZNU Faculty of Mathematics

Activities:
- M10–M15: Pilot delivery of Module 1 (AI/ML Fundamentals) to Cohort 1 (40 students). Pre-assessment of student competencies. Blended format: 60% online, 40% in-person laboratory.
- M13–M18: Deliver Modules 2 and 3 to Cohort 1. Summer school at NTNU for top 10 students (2 weeks).
- M18–M24: Full curriculum delivery to Cohort 2 (50 students). Module 4 piloted with Cohort 1. Student exchanges: 6 students to NTNU for semester-long study (up to 6 months).
- M24–M30: Full curriculum to Cohort 3 (60 students). Advanced projects in Living Lab. Student industry placements with Zaporizhzhia metallurgical companies.
- M30–M36: Final assessments, graduate tracking system setup, alumni network launch.

Deliverables: D3.1 Pilot course evaluation report (M18); D3.2 Student competency assessment results for all cohorts (M36); D3.3 Student exchange reports (M24, M36); D3.4 Graduate employment tracking framework (M36).

WP4: Living Laboratory and Digital Twin Pilots (Months 6–36)
Lead: ZNU Faculty of Mathematics | Co-lead: NTNU

Activities:
- M6–M12: Living Lab design and infrastructure setup. Connect ZNU Engineering Institute laboratories with Faculty of Mathematics computing resources. Install required software (Python, TensorFlow, simulation tools). Establish remote access to NTNU HPC resources for model training.
- M12–M18: Develop 2 educational digital twins: (i) electrolyser process for Mg extraction (based on ZNU research), (ii) alloy composition optimisation (using NTNU SFI datasets). Deploy on ZNU cloud infrastructure (Microsoft Azure).
- M18–M30: Student and staff use of Living Lab for coursework, research projects, and industry problem-solving. At least 3 industry-connected projects with Zaporizhzhia metallurgical companies.
- M30–M36: Living Lab sustainability plan: integration into ZNU's permanent research infrastructure, maintenance protocol, data management procedures.

Deliverables: D4.1 Living Lab technical specification and setup report (M12); D4.2 Two educational digital twins deployed (M18); D4.3 Industry collaboration reports (M30); D4.4 Living Lab sustainability plan (M36).

WP5: Dissemination and Sustainability (Months 1–36)
Lead: ZNU | Co-lead: NTNU

Activities:
- M1–M36: Project website and social media presence (LinkedIn, ResearchGate). Quarterly newsletter to stakeholders.
- M6, M18, M30: Annual project workshops (alternating ZNU/NTNU/online). Open to other Ukrainian HEIs and industry.
- M12–M36: Publication of 5+ joint papers in relevant journals (e.g., Journal of Materials Education, Computers & Education, Metallurgical and Materials Transactions).
- M18–M36: Development and publication of 4 Open Educational Resources (OER) on national and international repositories.
- M24–M36: Policy brief for Ukrainian Ministry of Education on AI-metallurgy curriculum integration. Replication guide for other Ukrainian technical universities.
- M30–M36: Preparation of follow-up proposals (Horizon Europe, Erasmus+ KA2) for continued cooperation.

Deliverables: D5.1 Dissemination plan and project website (M3); D5.2 Annual workshop reports (M6, M18, M30); D5.3 Joint publications (M36); D5.4 OER package (M36); D5.5 Policy brief and replication guide (M36); D5.6 Follow-up proposal drafts (M36).

Timeline overview:

Year 1 (M1–M12) — Foundation: Needs assessment, curriculum development, staff selection and first training round, Living Lab design, project website launch.
Year 2 (M13–M24) — Pilots: Pilot course delivery (Cohorts 1–2), staff exchanges (rounds 1–2), Living Lab operational with digital twins, first publications, first annual workshop.
Year 3 (M25–M36) — Scaling: Full curriculum delivery (Cohort 3), advanced Living Lab projects, industry collaboration, OER publication, sustainability planning, follow-up proposal preparation.

Teaching methodologies: The project employs problem-based learning (PBL) as its core pedagogical approach, complemented by flipped classroom methods (online theory + in-person practice), project-based industry collaborations, and peer learning across cohorts. The blended delivery model ensures resilience to security disruptions while maintaining quality standards equivalent to NTNU's own programmes."""

SECTION_2_2 = """The project has conducted a comprehensive risk assessment addressing all cross-cutting issues required by HK-dir. Each risk is rated by probability (Low/Medium/High) and impact (Low/Medium/High), with specific mitigation measures.

1. Security risks (Probability: High | Impact: High)
Zaporizhzhia is a frontline region directly affected by active hostilities. Risks include physical danger to participants, infrastructure damage, power outages, and disruption of educational activities.

Mitigation measures: (a) The blended delivery model is designed from the outset as a core feature, not an emergency fallback. All course modules have fully developed online versions accessible asynchronously. (b) ZNU has demonstrated operational resilience since 2022, maintaining educational delivery through air raid protocols, backup power systems, and Microsoft Azure cloud migration for critical IT systems. (c) Laboratory sessions will be scheduled during periods of relative stability, with alternative locations identified (Kyiv, Dnipro) for intensive workshops if Zaporizhzhia becomes inaccessible. (d) NTNU will provide remote access to computing resources, enabling AI model training regardless of local infrastructure status. (e) Emergency communication protocols and participant safety procedures aligned with Norwegian MFA travel advisories.

2. Corruption risk (Probability: Medium | Impact: High)
Ukraine's Corruption Perception Index score of 36/100 (Transparency International, 2023) triggers mandatory anti-corruption measures under Norwegian development cooperation guidelines [9].

Mitigation measures: (a) All procurement follows NTNU's established procedures, which comply with Norwegian Public Procurement Act and EEA regulations. (b) Dual-signature requirement for all financial transactions: both the NTNU project coordinator and ZNU project manager must approve expenditures above NOK 10,000. (c) Anti-corruption is a standing agenda item at all steering committee meetings. (d) Annual financial audits conducted by an independent auditor appointed by NTNU. (e) Transparent selection criteria for all scholarships, exchanges, and staff participation — published in advance with documented decision processes. (f) Whistleblower mechanism accessible to all project participants. (g) Anti-corruption training for all project staff during the kick-off phase (M1).

3. Climate and environmental risks (Probability: Low | Impact: Medium)
The project must minimise its carbon footprint while achieving educational objectives.

Mitigation measures: (a) Digital-first approach: all regular meetings, workshops, and training sessions use video conferencing as default. Physical travel is reserved for activities requiring hands-on laboratory work or in-person collaboration. (b) When physical travel is necessary, rail transport is preferred over air travel for European routes. (c) Staff exchanges are consolidated (longer stays, fewer trips) to reduce total travel emissions. (d) The project actively contributes to environmental sustainability through its core content: training engineers in green metallurgy and decarbonisation technologies that will reduce industrial emissions during Ukraine's reconstruction. (e) Carbon footprint tracked and reported annually using NTNU's institutional methodology.

4. Gender equality (Probability: Medium | Impact: Medium)
Metallurgical engineering is a male-dominated field in both Norway and Ukraine. Women are underrepresented among both students and academic staff in this discipline.

Mitigation measures: (a) Gender balance targets: aim for minimum 35% women among student participants and 40% among staff participants. (b) Active recruitment of women through targeted outreach to female students in engineering and mathematics programmes. (c) Gender-sensitive pedagogy: training for all teaching staff on inclusive classroom practices. (d) Course scheduling accommodates caregiving responsibilities (flexible deadlines, recorded sessions). (e) Female role models featured in guest lectures and case studies. (f) Gender-disaggregated data collected and reported for all project indicators. (g) ZNU's existing gender equality officer consulted on project design and implementation.

5. Human rights (Probability: Low | Impact: Medium)
The project must ensure non-discrimination, inclusive access, and attention to vulnerable groups, particularly in a conflict-affected context.

Mitigation measures: (a) Open and transparent selection processes for all student and staff participation. (b) Reasonable accommodations for participants with disabilities, including accessible online learning platforms. (c) Special attention to internally displaced persons (IDPs) among the student population — Zaporizhzhia hosts over 100,000 IDPs from occupied territories. (d) No discrimination based on gender, age, disability, ethnic origin, religion, or political views. (e) Participants' right to withdraw from any project activity without academic penalty. (f) Compliance with the Ukrainian Law on Higher Education regarding student rights and academic freedom.

6. SEAH prevention (Probability: Low | Impact: High)
Sexual exploitation, abuse, and harassment prevention is mandatory in all Norwegian-funded development cooperation.

Mitigation measures: (a) NTNU's institutional Code of Conduct on SEAH applies to all project activities and participants. (b) ZNU adopts or aligns its existing policies with Norwegian standards for SEAH prevention. (c) Mandatory SEAH awareness training for all project staff and exchange participants during orientation. (d) Confidential reporting mechanism established with designated contact persons at both NTNU and ZNU. (e) Clear procedures for investigation and follow-up of reported incidents, in line with Norwegian regulations. (f) Special attention during exchange visits when participants are in unfamiliar environments.

7. Data protection and IT security (Probability: Medium | Impact: Medium)
The project handles personal data of students and staff, and uses cloud-based computing resources.

Mitigation measures: (a) Compliance with GDPR (for NTNU) and Ukrainian data protection legislation (for ZNU). (b) Data processing agreements between partners before any personal data sharing. (c) ZNU's existing Microsoft Azure cloud infrastructure provides enterprise-grade security. (d) Research data stored on NTNU servers with institutional backup and access controls. (e) Student data anonymised in all publications and dissemination materials.

8. Academic freedom and national security (Probability: Low | Impact: Medium)
AI technologies have dual-use potential, and the project must be sensitive to export control considerations.

Mitigation measures: (a) Project focuses exclusively on educational applications of open-source AI/ML tools — no classified or restricted technologies involved. (b) All research datasets are civilian metallurgical data with no military application. (c) Export control screening conducted by NTNU's research administration for all technology transfer. (d) Academic freedom principles upheld: no censorship of research findings or teaching content.

9. Cost-effectiveness (cross-cutting)
The project budget of NOK 5,000,000 over 3 years represents strong cost-effectiveness given the scope of activities. Key justifications: (a) Norwegian personnel costs capped at 30% of budget, ensuring majority of funds flow to Ukrainian capacity building. (b) Blended delivery reduces travel costs while maintaining quality. (c) Use of open-source software (Python, TensorFlow, PyTorch) eliminates licensing costs. (d) Living Lab leverages existing ZNU laboratory infrastructure rather than building new facilities. (e) Curricula designed for permanence — once developed and accredited, no additional funding needed to deliver."""

SECTION_3_1 = """The AI-GreenMet consortium brings together complementary expertise from three institutional units at ZNU and NTNU, ensuring comprehensive coverage of all project domains without unnecessary overlap.

NTNU (Norwegian University of Science and Technology, Trondheim) — Coordinator

Project Coordinator: [To be named by NTNU]. The coordinator will be a senior faculty member from the Department of Materials Science and Engineering (IMA) with extensive experience in international educational cooperation and industry-academic partnerships. Required profile: PhD in metallurgy or materials science, minimum 10 years of university teaching experience, track record of EU/Nordic project coordination, and knowledge of AI applications in metallurgical processes.

Key NTNU personnel:
- Professor in Metallurgy/Materials Science (IMA): Expert in metal production, alloy design, and sustainable metallurgy. Contributes advanced metallurgical knowledge, access to SFI Metal Production research results, and supervision of curriculum content for Modules 3 (Green Metallurgy) and 4 (Quality Control). Estimated effort: 3 person-months/year.
- Associate Professor in AI/ML for Engineering (Department of Computer Science / IDI): Expert in machine learning applications for industrial processes, neural network architectures for materials property prediction. Leads development of Module 1 (AI/ML Fundamentals) and contributes to Module 2 (Digital Twins). Estimated effort: 2 person-months/year.
- Educational Developer (NTNU Teaching Excellence unit): Specialist in curriculum design, learning outcome formulation, and quality assurance. Supports pedagogical design of all modules and training of ZNU staff in modern teaching methods. Estimated effort: 1 person-month/year.
- Project Administrator: Handles financial management, reporting, compliance, and exchange logistics. Estimated effort: 2 person-months/year.

Zaporizhzhia National University (ZNU) — Main Ukrainian Partner

ZNU Project Manager: A senior academic with cross-departmental authority, responsible for local coordination across ZNU's three participating units. Required profile: experience in international project management, English proficiency, knowledge of Ukrainian higher education regulations for curriculum changes.

ZNU Engineering Educational and Scientific Institute:
- Head of Metallurgy Department: Expert in high-entropy alloys, thermochemical processing, and industrial metallurgy. Leads ZNU-side curriculum development for metallurgical content, coordinates industry partnerships with Zaporizhzhia metallurgical enterprises, and manages the Living Lab's metallurgical equipment. Estimated effort: 4 person-months/year.
- Senior Researcher in Electrochemistry: Specialist in electrolyser systems for magnesium extraction from fluoride-oxide melts. Contributes real research data and case studies for the educational digital twin development. Estimated effort: 2 person-months/year.
- Researcher in Catalysis: Expert in catalytic technologies for industrial gas purification. Provides additional metallurgical case studies and contributes to Module 3 content. Estimated effort: 1 person-month/year.

ZNU Faculty of Mathematics:
- Head of Laboratory of Parallel and Distributed Computing: Expert in AI/ML model development, high-performance computing, and parallel algorithms. Leads the AI/computing side of curriculum development, manages the Living Lab's computing infrastructure, and supervises student AI projects. Estimated effort: 4 person-months/year.
- Associate Professor in Data Science: Specialist in data-driven optimisation and machine learning for engineering applications. Co-develops Modules 1 and 2, teaches programming components (Python, TensorFlow/PyTorch). Estimated effort: 3 person-months/year.

ZNU Metallurgical Professional College:
- Deputy Director for Academic Affairs: Coordinates vocational-level adaptation of course materials for technician students. Ensures College students can access simplified versions of AI-metallurgy modules. Estimated effort: 2 person-months/year.

[Partner 2 — To be confirmed by NTNU]:
The second Ukrainian HEI will contribute complementary metallurgical or AI expertise, expanding the student base and ensuring geographic diversity beyond Zaporizhzhia. Key personnel will be identified upon partner confirmation.

Complementarity: NTNU provides world-class metallurgy research, AI expertise, and educational quality assurance. ZNU Engineering Institute contributes domain-specific metallurgical research and industry connections. ZNU Faculty of Mathematics provides AI/ML and computing infrastructure. ZNU Metallurgical College ensures the project reaches vocational education. This multi-unit approach creates a pipeline from vocational to bachelor's to master's level, with no significant overlap between partners' contributions."""

SECTION_3_2 = """Division of responsibilities:

NTNU (Coordinator — ~30% of budget):
- Overall project coordination, financial management, and reporting to HK-dir
- Leadership of WP1 (Curriculum Development) and WP2 (Staff Training)
- Quality assurance of all educational content against European standards
- Hosting staff and student exchanges in Trondheim
- Providing remote access to computing resources and research datasets
- Leading preparation of follow-up proposals (Horizon Europe, Erasmus+)

ZNU (Main Ukrainian Partner — ~55% of budget):
- Local project management and coordination across three ZNU units
- Leadership of WP3 (Student Training) and WP5 (Dissemination & Sustainability)
- Co-leadership of WP4 (Living Lab) — hosting physical infrastructure
- Student recruitment, course delivery, and assessment
- Industry liaison with Zaporizhzhia metallurgical enterprises
- Integration of new curricula into ZNU's official degree programme catalogue
- Ensuring compliance with Ukrainian higher education regulations

[Partner 2 — TBD] (~15% of budget):
- Participation in curriculum development for adaptation to their institutional context
- Delivery of adapted course modules to their own students (expanding geographic reach)
- Contribution of complementary expertise (metallurgical or AI, depending on partner)
- Participation in dissemination activities and policy dialogue

Management structure:

(a) Steering Committee: Meets quarterly (online, with one annual in-person meeting). Comprises: NTNU Project Coordinator (chair), ZNU Project Manager, Partner 2 representative, one external advisor from Norwegian-Ukrainian cooperation framework. Responsibilities: strategic oversight, approval of annual work plans and budgets, risk management, conflict resolution.

(b) Project Management Team: Meets monthly (online). Comprises: NTNU Project Coordinator, ZNU Project Manager, WP leaders, NTNU project administrator. Responsibilities: operational coordination, monitoring of deliverables and timeline, financial tracking, preparation of steering committee meetings.

(c) Work Package Leaders: Each WP has a designated leader and co-leader from different partner institutions, ensuring cross-institutional coordination. WP leaders report monthly to the Project Management Team.

Communication and coordination mechanisms:
- Microsoft Teams workspace for daily communication, document sharing, and file management
- Shared project management tool (e.g., Asana or Trello) for task tracking and deadline management
- Quarterly progress reports prepared by WP leaders
- Annual in-person project workshop (rotating: Year 1 at NTNU, Year 2 at ZNU, Year 3 at Partner 2 or online)
- Shared document repository on NTNU's institutional cloud storage (OneDrive/SharePoint) with access for all partners

Stakeholder engagement plan:
- Zaporizhzhia metallurgical industry: Advisory role in curriculum needs assessment (WP1), guest lectures, student industry projects (WP3), Living Lab industry collaborations (WP4). Key contacts: Zaporizhstal (steel production), regional metallurgical SMEs.
- Ukrainian Ministry of Education and Science: Consulted on curriculum accreditation requirements, invited to annual workshops, recipient of policy brief (WP5).
- Norwegian-Ukrainian cooperation networks: HK-dir programme community, Norwegian Embassy in Ukraine, NTNU International Office.
- Other Ukrainian HEIs: Invited to annual workshops, recipients of replication guide and OER materials.
- European networks: EUA (European University Association), SEFI (European Society for Engineering Education) — for dissemination and benchmarking."""

SECTION_4_1 = """The AI-GreenMet project is designed to create measurable, multi-level impact extending from individual students to the national education system and regional industrial ecosystem.

Direct impact on target groups:

(a) Students (150+ over 3 years): Graduates will possess a unique interdisciplinary skill set combining metallurgical engineering with AI/ML competencies — a profile currently unavailable from any Ukrainian HEI. Specific competencies acquired: programming in Python for engineering applications; application of machine learning to metallurgical process data; design and use of digital twins; understanding of green metallurgy and decarbonisation technologies; data-driven quality control methods. Pre/post competency assessments (WP3) will document measurable learning gains. Employment prospects are significantly enhanced: reconstructed metallurgical enterprises in Zaporizhzhia and beyond require precisely these interdisciplinary skills, and European employers increasingly seek engineers with AI competencies.

(b) Academic staff (25+ trained): Participating staff will gain new pedagogical skills (problem-based learning, blended delivery, digital twin pedagogy) and technical competencies (AI/ML tools for metallurgical applications). This capacity remains at ZNU permanently after the project ends, enabling continued delivery and further development of the curriculum. Each trained staff member can reach 50–100 students per year, creating a multiplier effect.

(c) Metallurgical Professional College students (40+): Vocational students gain practical AI-enhanced skills (basic data analysis, digital instrument operation, quality monitoring using AI tools) applicable to technician-level roles in modern metallurgical facilities.

(d) Industry professionals (50+ through short courses and Living Lab): The Zaporizhzhia metallurgical cluster gains access to continuing education and problem-solving support through the Living Lab, facilitating technology adoption during reconstruction.

Institutional impact:

(a) ZNU: The project establishes AI-metallurgy as a permanent educational specialisation within the Engineering Institute, strengthens research capacity through the Living Lab and NTNU partnership, and positions ZNU as a national leader in green metallurgy education. The Faculty of Mathematics gains a new application domain for its AI/ML research, increasing its relevance and funding competitiveness.

(b) NTNU: Gains a strategic Ukrainian partner for educational cooperation, access to unique metallurgical datasets and research challenges, and enhanced capacity for international development projects. The partnership supports NTNU's own strategic priorities in sustainable industry and global engagement.

(c) ZNU Metallurgical College: Connects to university-level curriculum development for the first time, creating a clear educational pathway from vocational to bachelor's to master's level in AI-enhanced metallurgy.

Regional and national impact:

(a) Zaporizhzhia region: The metallurgical cluster (historically employing 50,000+ workers) requires engineers with modern skills for reconstruction. AI-GreenMet directly addresses this workforce need. The Living Lab creates a permanent resource for industry-academic collaboration, supporting technology transfer and innovation during the region's post-war economic recovery.

(b) National scale: The curriculum modules, teaching materials, and OER produced by the project are designed for replication by other Ukrainian technical universities — particularly the National Metallurgical Academy of Ukraine (Dnipro), National Technical University "Kharkiv Polytechnic Institute", and KPI Kyiv. The policy brief (WP5) will propose integration of AI-metallurgy education into national STEM education standards.

Dissemination plan:

(1) Scientific publications: 5+ joint papers in peer-reviewed journals (target outlets: Journal of Materials Education, International Journal of Engineering Education, Computers & Education, Metallurgical and Materials Transactions B). All published as open access.

(2) Conference presentations: Annual presentations at relevant conferences: SEFI (European Society for Engineering Education), TMS (The Minerals, Metals & Materials Society), and Ukrainian national education conferences.

(3) Open Educational Resources: 4 OER packages (one per curriculum module) published on NTNU's Open Learning repository and the Ukrainian national OER platform. Each package includes: syllabus, lecture materials, laboratory exercises, programming assignments, assessment rubrics.

(4) Project website and social media: Maintained throughout the project and for 2 years after (ZNU responsibility). Content: news, publications, OER downloads, student testimonials, industry case studies.

(5) Policy brief: Distributed to the Ukrainian Ministry of Education and Science, Ukrainian higher education rectors' councils, and Norwegian-Ukrainian cooperation stakeholders. Proposes a national framework for AI-enhanced engineering education.

(6) Replication guide: Step-by-step guide for other Ukrainian HEIs to establish similar AI-metallurgy programmes, including curriculum templates, technology requirements, partnership models, and lessons learned.

Impact beyond the consortium: The project addresses a systemic gap in Ukrainian higher education. By making all educational materials openly available and providing a replication guide, the impact extends to the estimated 15 Ukrainian universities offering metallurgy-related programmes. The Living Lab model is transferable to other industrial sectors (chemical engineering, energy systems, manufacturing) and regions."""

SECTION_4_2 = """Sustainability plan — ensuring impact beyond the project period:

(a) Curriculum integration: The 4 course modules (20 ECTS total) will be formally accredited and integrated into ZNU's Engineering Institute degree programmes (bachelor's in Metallurgical Engineering, master's in Materials Science) by Month 24. Once embedded in the official curriculum, delivery requires no additional external funding — it becomes part of ZNU's regular teaching load. Ukrainian higher education law guarantees that accredited programmes continue as long as student demand exists.

(b) Living Lab as permanent infrastructure: The Living Lab established in WP4 will be integrated into ZNU's institutional research infrastructure. Operating costs (electricity, computing, maintenance) are covered by ZNU's regular budget. The Lab will generate its own income through industry consulting projects with Zaporizhzhia metallurgical enterprises and externally funded research grants.

(c) Long-term partnership: The NTNU-ZNU cooperation is not limited to this project. The partners have agreed on a joint roadmap extending to 2028 and beyond: (i) 2026–2028: AI-GreenMet project activities; (ii) 2028+: co-supervised PhD projects in AI-metallurgy, joint Horizon Europe proposals (targeting MSCA Doctoral Networks and Cluster 4 — Digital, Industry and Space), continued staff exchanges through Erasmus+ ICM. The Steering Committee will approve a formal partnership extension agreement by Month 30.

(d) Trained staff multiplier: The 25+ trained ZNU staff members will continue teaching the new curriculum modules and training their colleagues after the project ends. ZNU will establish an internal mentoring programme where project-trained staff guide new faculty in AI-enhanced metallurgical education, creating a self-sustaining capacity development mechanism.

(e) OER and knowledge base: All Open Educational Resources remain available on public repositories indefinitely. Other Ukrainian universities can adopt and adapt these materials without project involvement.

Evaluation methodology:

(1) Pre/post competency assessments: All participating students complete a standardised assessment at the start and end of each module, measuring knowledge and practical skills in AI/ML and metallurgy. Instruments developed collaboratively by NTNU and ZNU based on established engineering education assessment frameworks [10]. Results compared against a control group of non-participating engineering students.

(2) Staff competency tracking: Teaching staff complete self-assessment surveys before and after training activities, measuring confidence and skill in: AI/ML tools, digital twin development, problem-based learning pedagogy, and blended course delivery.

(3) Student satisfaction surveys: Anonymous surveys after each module, using a standardised questionnaire (adapted from NTNU's institutional evaluation system) covering: course quality, teaching effectiveness, relevance to career goals, learning environment.

(4) Graduate employment tracking: ZNU's career centre will track graduates of the AI-metallurgy programme for 3 years after completion, recording: employment rates, employer satisfaction, relevance of AI/ML skills to job roles, salary levels compared to graduates without AI training.

(5) Industry satisfaction surveys: Annual surveys of participating industry partners (Zaporizhzhia metallurgical enterprises) measuring: relevance of graduate skills to industry needs, quality of Living Lab collaborations, willingness to continue engagement.

(6) Annual external review: An independent evaluator (appointed by the Steering Committee, not a project partner) conducts annual reviews assessing progress against planned results, quality of outputs, and strategic alignment. Review reports submitted to HK-dir with annual progress reports.

Documentation and reporting:
- Annual progress reports to HK-dir including financial statements and activity documentation
- Mid-term evaluation report (M18) with recommendations for Year 3 adjustments
- Final project report (M36) with comprehensive impact assessment, lessons learned, and sustainability documentation
- All evaluation data integrated into ZNU's institutional management information system for long-term tracking
- Public dissemination of aggregated evaluation results through publications and OER repository"""

REFERENCES = [
    "[1] KSE Institute (2024). Russia Will Pay: Assessing Damages to Ukraine's Infrastructure. Kyiv School of Economics, Report No. 7.",
    "[2] European Commission (2023). Towards Competitive and Clean European Steel — Communication from the Commission. COM(2023) 219 final.",
    "[3] NTNU (2024). SFI Metal Production: Centre for Research-based Innovation. Annual Report 2023. Norwegian University of Science and Technology, Trondheim.",
    "[4] HK-dir (2025). Nansen EDU — Call for Applications 2025: Programme Document. Directorate for Higher Education and Skills, Norway.",
    "[5] Zaporizhzhia National University (2024). Research Activity Report 2023. ZNU, Zaporizhzhia, Ukraine.",
    "[6] European Commission (2024). EU4Environment — Green Economy in the Eastern Partnership. Programme Progress Report.",
    "[7] Malmqvist, J., Rådberg, K.K., & Lundqvist, U. (2015). Comparative Analysis of Challenge-Based Learning Experiences. Proc. 11th International CDIO Conference, Chengdu, China.",
    "[8] World Steel Association (2023). Steel Industry and the Green Deal: Decarbonisation Pathways for the European Steel Sector. Brussels.",
    "[9] Transparency International (2023). Corruption Perception Index 2023. Berlin.",
    "[10] Crawley, E.F., Malmqvist, J., Östlund, S., Brodeur, D.R., & Edström, K. (2014). Rethinking Engineering Education: The CDIO Approach. 2nd ed. Springer.",
    "[11] Ministry of Education and Science of Ukraine (2023). Strategy for the Development of Higher Education in Ukraine 2022–2032. Kyiv.",
    "[12] OECD (2023). Rebuilding Ukraine: Principles for Reconstruction. OECD Publishing, Paris.",
    "[13] Winkler-Schwartz, A., et al. (2022). Artificial Intelligence in Engineering Education: A Systematic Review. Journal of Engineering Education, 111(4), 745–767.",
    "[14] European Parliament (2024). Carbon Border Adjustment Mechanism (CBAM): Implementation Regulation. Regulation (EU) 2023/956.",
    "[15] Norwegian Ministry of Foreign Affairs (2023). Norway's Support for Ukraine: Nansen Programme. Oslo."
]


# ============================================================
# FILL THE DOCX TEMPLATE
# ============================================================

def fill_template():
    doc = Document(TEMPLATE_PATH)

    # --- Fill Cover Table (Table 0) ---
    cover_table = doc.tables[0]
    cover_table.rows[1].cells[1].text = "AI-Driven Green Metallurgy Education for Ukraine's Industrial Reconstruction (AI-GreenMet)"
    cover_table.rows[2].cells[1].text = "NTNU — Norwegian University of Science and Technology"
    # Project number left blank (assigned by Espresso)

    # --- Delete all guideline tables (Tables 1–14) ---
    # We need to remove them from the XML. Do it in reverse order to avoid index shifts.
    tables_to_delete = list(range(1, 15))  # Tables 1 through 14
    for idx in reversed(tables_to_delete):
        table = doc.tables[idx]
        table_element = table._element
        table_element.getparent().remove(table_element)

    # --- Map section content to paragraph indices ---
    # Find the "► Delete the guidelines boxes and write here." paragraphs and replace them
    sections = {
        '1.1': SECTION_1_1,
        '1.2': SECTION_1_2,
        '2.1': SECTION_2_1,
        '2.2': SECTION_2_2,
        '3.1': SECTION_3_1,
        '3.2': SECTION_3_2,
        '4.1': SECTION_4_1,
        '4.2': SECTION_4_2,
    }

    # Find section headers and their corresponding "write here" markers
    current_section = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Track which section we're in
        if text.startswith('1.1'):
            current_section = '1.1'
        elif text.startswith('1.2'):
            current_section = '1.2'
        elif text.startswith('2.1'):
            current_section = '2.1'
        elif text.startswith('2.2'):
            current_section = '2.2'
        elif text.startswith('3.1'):
            current_section = '3.1'
        elif text.startswith('3.2'):
            current_section = '3.2'
        elif text.startswith('4.1'):
            current_section = '4.1'
        elif text.startswith('4.2'):
            current_section = '4.2'
        elif '► Delete' in text and current_section in sections:
            # Replace this paragraph with section content
            content = sections[current_section]
            paragraphs_text = [p.strip() for p in content.strip().split('\n\n') if p.strip()]

            # Set first paragraph in place of the marker
            para.text = paragraphs_text[0]
            para.style = doc.styles['Normal']
            # Set font
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Insert remaining paragraphs after the first one
            # We need to insert new paragraphs after the current one in the XML
            current_element = para._element
            for ptext in paragraphs_text[1:]:
                new_para = doc.add_paragraph()
                new_para.text = ptext
                new_para.style = doc.styles['Normal']
                for run in new_para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                # Move the new paragraph right after current_element
                current_element.addnext(new_para._element)
                current_element = new_para._element

            current_section = None  # Done with this section

    # --- Fill References Section ---
    # Find the [Reference] placeholders and replace them
    ref_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in ('[Reference]', '[...]'):
            ref_indices.append(i)

    # Remove old reference placeholders and add new ones
    if ref_indices:
        # Replace first [Reference] with first actual reference
        first_ref_para = doc.paragraphs[ref_indices[0]]
        first_ref_para.text = REFERENCES[0]
        first_ref_para.style = doc.styles['Normal']
        for run in first_ref_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

        # Add remaining references after the first one
        current_element = first_ref_para._element
        for ref in REFERENCES[1:]:
            new_para = doc.add_paragraph()
            new_para.text = ref
            new_para.style = doc.styles['Normal']
            for run in new_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            current_element.addnext(new_para._element)
            current_element = new_para._element

        # Remove remaining old [Reference] and [...] placeholders
        for idx in ref_indices[1:]:
            para = doc.paragraphs[idx]
            if para.text.strip() in ('[Reference]', '[...]'):
                para.text = ''

    # --- Remove "List of references referred to..." line ---
    for para in doc.paragraphs:
        if 'List of references referred to in the project description' in para.text:
            para.text = ''

    # --- Clean up empty paragraphs around sections ---
    # (Keep minimal spacing)

    # Save
    doc.save(OUTPUT_PATH)
    print(f'Template filled and saved to: {OUTPUT_PATH}')

    # Print statistics
    total_words = 0
    section_names = {
        '1.1': 'Background and needs',
        '1.2': 'Alignment and innovation',
        '2.1': 'Activities and work plan',
        '2.2': 'Risk factors',
        '3.1': 'Key personnel',
        '3.2': 'Collaboration',
        '4.1': 'Impact and dissemination',
        '4.2': 'Sustainability and evaluation',
    }
    all_sections = [
        ('1.1', SECTION_1_1), ('1.2', SECTION_1_2),
        ('2.1', SECTION_2_1), ('2.2', SECTION_2_2),
        ('3.1', SECTION_3_1), ('3.2', SECTION_3_2),
        ('4.1', SECTION_4_1), ('4.2', SECTION_4_2),
    ]
    for key, text in all_sections:
        words = len(text.split())
        total_words += words
        print(f'  Section {key} ({section_names[key]}): ~{words} words')

    print(f'\n  References: {len(REFERENCES)}')
    print(f'  Total words: ~{total_words}')
    print(f'  Estimated pages (11pt Arial): ~{round(total_words / 400)}')


if __name__ == '__main__':
    fill_template()
